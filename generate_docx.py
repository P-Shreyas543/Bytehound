"""
Convert app/resources/index.html to a formatted DOCX (Bytehound User Manual).
Run: python generate_docx.py
Output: Documentation/Bytehound_User_Manual.docx
taskkill /F /IM winword.exe ; .venv\Scripts\python.exe generate_docx.py
"""

import json
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR       = Path(__file__).resolve().parent
HTML_PATH      = BASE_DIR / "app" / "resources" / "index.html"
VERSION_PATH   = BASE_DIR / "version.json"
TEMPLATE_PATH  = BASE_DIR / "Documentation" / "Template.docx"
COVER_META_PATH = BASE_DIR / "cover_meta.json"
OUT_PATH       = BASE_DIR / "Documentation" / "Bytehound_User_Manual.docx"

# ── Brand colours ──────────────────────────────────────────────────────────────
# Matches the index.html colour palette exactly
BRAND_DARK    = RGBColor(0x1A, 0x5F, 0x7A)   # #1a5f7a — headings / brand
BRAND_ACCENT  = RGBColor(0x34, 0x98, 0xDB)   # #3498db — accent / rule lines
TEXT_BODY     = RGBColor(0x2C, 0x3E, 0x50)   # #2c3e50 — body text
CODE_FG       = RGBColor(0x1A, 0x5F, 0x7A)   # inline code foreground
FONT_NAME     = "PT Sans"
FONT_MONO     = "Consolas"

# Callout palette: border_hex, bg_hex, label
_CALLOUT_STYLES = {
    "note":    ("2196F3", "E3F2FD", "ℹ  Note"),
    "warning": ("FFC107", "FFF8E1", "⚠  Warning"),
    "success": ("4CAF50", "E8F5E9", "✔  Tip"),
}


# ── XML / table helpers ────────────────────────────────────────────────────────
def _get_or_add_tblPr(tbl_element):
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_element.insert(0, tblPr)
    return tblPr


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_no_borders(table):
    tblPr = _get_or_add_tblPr(table._tbl)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def set_cell_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "right", "bottom", "left"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_grid_borders(table, color="C0C8D8"):
    tblPr = _get_or_add_tblPr(table._tbl)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def set_cell_left_border(cell, color_hex: str, sz: str = "24"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "right", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color_hex)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def add_bottom_border_to_para(para, color_hex: str, sz: str = "4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Run helpers ────────────────────────────────────────────────────────────────
def apply_run_font(run, size=Pt(11), bold=False, italic=False,
                   color: RGBColor | None = None, mono: bool = False):
    run.font.name = FONT_MONO if mono else FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_inline_runs(para, node, bold=False, italic=False, size=Pt(11)):
    """Recursively turn inline HTML nodes into Word runs on *para*."""
    if isinstance(node, NavigableString):
        text = str(node)
        text = re.sub(r'\s+', ' ', text)
        if text == '' or text == ' ':
            if len(para.runs) == 0:
                return
        if len(para.runs) == 0:
            text = text.lstrip()
        if text:
            run = para.add_run(text)
            apply_run_font(run, size=size, bold=bold, italic=italic)
        return

    tag = getattr(node, "name", None)
    if tag is None:
        return

    # Propagate formatting down
    is_bold   = bold   or tag in ("strong", "b")
    is_italic = italic or tag in ("em", "i")
    is_code   = tag in ("code", "kbd")
    is_link   = tag == "a"

    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            text = re.sub(r'\s+', ' ', text)
            if text == '' or text == ' ':
                if len(para.runs) == 0:
                    continue
            if len(para.runs) == 0:
                text = text.lstrip()
            if text:
                run = para.add_run(text)
                if is_code:
                    apply_run_font(run, size=Pt(10), bold=is_bold, italic=is_italic,
                                   mono=True, color=CODE_FG)
                elif is_link:
                    apply_run_font(run, size=size, bold=is_bold, italic=is_italic,
                                   color=BRAND_ACCENT)
                    run.font.underline = True
                else:
                    apply_run_font(run, size=size, bold=is_bold, italic=is_italic)
        elif hasattr(child, "name"):
            add_inline_runs(para, child, bold=is_bold, italic=is_italic, size=size)


def strip_para_ends(para):
    """Strip leading and trailing whitespace from the runs of a paragraph."""
    if para.runs:
        para.runs[0].text = para.runs[0].text.lstrip()
        para.runs[-1].text = para.runs[-1].text.rstrip()


# ── Callout blocks (note / warning / success) ──────────────────────────────────
def add_callout(doc, element):
    cls_list = element.get("class", [])
    cls_str = " ".join(cls_list) if isinstance(cls_list, list) else cls_list
    style_key = next((k for k in _CALLOUT_STYLES if k in cls_str), None)
    if style_key is None:
        # Unknown div — recurse into its children
        for child in element.children:
            process_element(doc, child)
        return

    border_hex, bg_hex, label = _CALLOUT_STYLES[style_key]

    # 1-column, 1-row table to fake left-border + background
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Normal Table"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(tbl)

    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    set_cell_left_border(cell, border_hex)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(6)

    label_run = p.add_run(label + "  ")
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.size = Pt(10)

    # Gather content from all children (skip nested headings)
    for child in element.children:
        if hasattr(child, "name") and child.name in ("h3", "h4"):
            heading_text = child.get_text(strip=True)
            hr = p.add_run(heading_text + "  ")
            hr.bold = True
            hr.font.name = FONT_NAME
            hr.font.size = Pt(10)
        else:
            add_inline_runs(p, child, size=Pt(10))

    strip_para_ends(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Pre / code-block ───────────────────────────────────────────────────────────
def add_code_block(doc, text: str):
    """Render a <pre> block as a shaded monospace table cell."""
    # Strip leading/trailing blank lines
    lines = text.rstrip("\n")

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Normal Table"
    set_table_no_borders(tbl)

    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F0F0F0")
    set_cell_no_borders(cell)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(8)

    run = p.add_run(lines)
    run.font.name = FONT_MONO
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── HTML table → Word table ───────────────────────────────────────────────────
def add_html_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return

    max_cols = max(
        sum(int(td.get("colspan", 1)) for td in row.find_all(["th", "td"]))
        for row in rows
    )
    if max_cols == 0:
        return

    wtable = doc.add_table(rows=len(rows), cols=max_cols)
    wtable.style = "Normal Table"
    wtable.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_grid_borders(wtable)

    for r_idx, row in enumerate(rows):
        wrow = wtable.rows[r_idx]

        # Prevent any single row from splitting across a page break
        trPr = wrow._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        cant.set(qn("w:val"), "1")
        trPr.append(cant)

        cells = row.find_all(["th", "td"])
        c_idx = 0
        for cell_tag in cells:
            if c_idx >= max_cols:
                break
            wcell = wrow.cells[c_idx]

            is_header = cell_tag.name == "th" or r_idx == 0
            if is_header:
                set_cell_bg(wcell, "1A5F7A")
            elif r_idx % 2 == 1:
                set_cell_bg(wcell, "F4F6F8")
            else:
                set_cell_bg(wcell, "FFFFFF")

            wcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = wcell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Pt(4)

            # Header row: keepNext ensures it stays on the same page as row below it
            if is_header:
                pPr = para._p.get_or_add_pPr()
                kn = OxmlElement("w:keepNext")
                kn.set(qn("w:val"), "1")
                pPr.append(kn)

            txt_color = RGBColor(0xFF, 0xFF, 0xFF) if is_header else RGBColor(0x33, 0x33, 0x33)
            add_inline_runs(para, cell_tag, bold=is_header, size=Pt(10))
            for r in para.runs:
                r.font.color.rgb = txt_color

            strip_para_ends(para)
            c_idx += int(cell_tag.get("colspan", 1))

def add_image_element(doc, img_tag):
    """Render an <img> tag as a centered image with an optional caption."""
    src = img_tag.get("src")
    if not src:
        return
    
    # Resolve relative paths
    img_path = Path(src)
    if not img_path.is_absolute():
        # Relative to app/resources
        img_path = BASE_DIR / "app" / "resources" / img_path

    if not img_path.exists():
        print(f"  Warning: Image file not found: {img_path}")
        return

    # Add picture (width 14cm is a safe choice to fit inside the margins)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    
    # Set keep_with_next if possible via pPr
    pPr = p._p.get_or_add_pPr()
    kn = OxmlElement("w:keepNext")
    kn.set(qn("w:val"), "1")
    pPr.append(kn)
    
    # Determine layout width based on the CSS class from the HTML
    cls_list = img_tag.get("class", [])
    if isinstance(cls_list, str):
        cls_list = cls_list.split()
    
    img_width = Cm(14)  # Default full page width
    if "img-small" in cls_list:
        img_width = Cm(7)
    elif "img-medium" in cls_list:
        img_width = Cm(11)

    run = p.add_run()
    run.add_picture(str(img_path), width=img_width)

    # Alt text as caption below the image
    alt = img_tag.get("alt")
    if alt:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(12)
        
        run_cap = p_cap.add_run(f"Figure: {alt}")
        run_cap.italic = True
        run_cap.font.name = FONT_NAME
        run_cap.font.size = Pt(9.5)
        run_cap.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


# ── List handling ──────────────────────────────────────────────────────────────
def add_list_items(doc, ul_tag, ordered=False, depth=0):
    for idx, li in enumerate(ul_tag.find_all("li", recursive=False), start=1):
        nested = []
        text_parts = []
        for child in li.children:
            if hasattr(child, "name") and child.name in ("ul", "ol"):
                nested.append(child)
            else:
                text_parts.append(child)

        bullet = f"{idx}." if ordered else "•"
        indent = Pt(18 * (depth + 1))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent       = indent + Pt(14)
        p.paragraph_format.first_line_indent = Pt(-14)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(1)

        br = p.add_run(bullet + "  ")
        apply_run_font(br, size=Pt(11), bold=True,
                       color=BRAND_ACCENT if not ordered else TEXT_BODY)

        for part in text_parts:
            add_inline_runs(p, part, size=Pt(11))

        strip_para_ends(p)

        for nested_list in nested:
            add_list_items(doc, nested_list,
                           ordered=(nested_list.name == "ol"),
                           depth=depth + 1)


# ── TOC field (Word will populate on Ctrl+A → F9) ────────────────────────────
def insert_toc_field(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = para.add_run("Right-click → 'Update Field' to populate the Table of Contents.")
    run4.font.name = FONT_NAME
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run4.font.italic = True

    run5 = para.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)


# ── Page setup (A4, standard margins) ─────────────────────────────────────────
def setup_page(doc):
    sec = doc.sections[0]
    sec.page_height    = Cm(29.7)
    sec.page_width     = Cm(21.0)
    sec.left_margin    = Cm(2.54)
    sec.right_margin   = Cm(2.54)
    sec.top_margin     = Cm(2.54)
    sec.bottom_margin  = Cm(2.0)


# ── Cover page ─────────────────────────────────────────────────────────────────
def build_cover_page(doc, version_info: dict):
    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(60)
    p_title.paragraph_format.space_after  = Pt(6)
    r = p_title.add_run("Bytehound")
    r.font.name  = FONT_NAME
    r.font.size  = Pt(40)
    r.bold       = True
    r.font.color.rgb = BRAND_DARK

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(6)
    r2 = p_sub.add_run("Serial Telemetry & Decoder Guide")
    r2.font.name  = FONT_NAME
    r2.font.size  = Pt(20)
    r2.font.color.rgb = BRAND_ACCENT

    # Tagline
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_after = Pt(30)
    r3 = p_tag.add_run(
        "Framed UART & Modbus RTU telemetry · live plots · TX commands · offline analysis"
    )
    r3.font.name   = FONT_NAME
    r3.font.size   = Pt(11)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # Horizontal rule
    hr = doc.add_paragraph()
    add_bottom_border_to_para(hr, "3498DB", sz="8")
    hr.paragraph_format.space_after = Pt(20)

    # Metadata info table — rows driven entirely by cover_meta.json
    # Each row: { "label": "...", "version_key": "<key in version.json> | __NUMPAGES__" }
    cover_meta: list = []
    if COVER_META_PATH.exists():
        with open(COVER_META_PATH, encoding="utf-8") as _f:
            cover_meta = json.load(_f).get("table", [])

    meta = []   # list of (label, value_or_token)
    for row in cover_meta:
        label = row.get("label", "")
        key   = row.get("version_key", "").strip()
        if key == "__NUMPAGES__":
            value = "__NUMPAGES__"          # rendered as a Word field below
        else:
            value = version_info.get(key, "") if key else row.get("value", "")
        meta.append((label, value))

    info_tbl = doc.add_table(rows=len(meta), cols=2)
    info_tbl.style = "Normal Table"
    set_table_grid_borders(info_tbl, "AAAAAA")

    # Centre the table and set fixed width
    tblPr = _get_or_add_tblPr(info_tbl._tbl)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "8000")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    for i, (label, val) in enumerate(meta):
        row = info_tbl.rows[i]
        set_cell_bg(row.cells[0], "1A5F7A")
        set_cell_bg(row.cells[1], "F4F6F8" if i % 2 == 0 else "FFFFFF")

        # Label cell (left, white text)
        lbl_cell = row.cells[0]
        lpp = lbl_cell.paragraphs[0]
        lpp.paragraph_format.space_before = Pt(6)
        lpp.paragraph_format.space_after  = Pt(6)
        lpp.paragraph_format.left_indent  = Pt(8)
        lr = lpp.add_run(label)
        lr.font.name      = FONT_NAME
        lr.font.size      = Pt(11)
        lr.bold           = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Value cell (right)
        val_cell = row.cells[1]
        vpp = val_cell.paragraphs[0]
        vpp.paragraph_format.space_before = Pt(6)
        vpp.paragraph_format.space_after  = Pt(6)
        vpp.paragraph_format.left_indent  = Pt(8)

        if val == "__NUMPAGES__":
            # Insert a { NUMPAGES } auto-count field
            def _field_run(para, fld_type=None, instr=None, placeholder=None):
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22")
                rPr.append(sz)
                r.append(rPr)
                if fld_type:
                    fc = OxmlElement("w:fldChar")
                    fc.set(qn("w:fldCharType"), fld_type)
                    r.append(fc)
                if instr:
                    it = OxmlElement("w:instrText")
                    it.text = instr
                    r.append(it)
                if placeholder is not None:
                    t = OxmlElement("w:t"); t.text = str(placeholder)
                    r.append(t)
                para._p.append(r)
            _field_run(vpp, fld_type="begin")
            _field_run(vpp, instr="NUMPAGES")
            _field_run(vpp, fld_type="separate")
            _field_run(vpp, placeholder="?")
            _field_run(vpp, fld_type="end")
        else:
            vr = vpp.add_run(val)
            vr.font.name      = FONT_NAME
            vr.font.size      = Pt(11)
            vr.font.color.rgb = BRAND_DARK

    # Table of Contents section
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    toc_hdr = doc.add_paragraph()
    toc_hdr.paragraph_format.space_before = Pt(24)
    toc_hdr.paragraph_format.space_after  = Pt(8)
    r4 = toc_hdr.add_run("Table of Contents")
    r4.font.name  = FONT_NAME
    r4.font.size  = Pt(16)
    r4.bold       = True
    r4.font.color.rgb = BRAND_DARK
    add_bottom_border_to_para(toc_hdr, "3498DB", sz="4")

    insert_toc_field(doc)
    doc.add_page_break()


# ── Heading renderer ──────────────────────────────────────────────────────────
# HTML h2 → Word "Heading 1" (main sections, numbered 1-7 in the manual)
# HTML h3 → Word "Heading 2" (subsections)
# HTML h4 → Word "Heading 3" (sub-subsections)
_HEADING_MAP = {
    "h2": ("Heading 1", Pt(18), Pt(18), Pt(8),  True),
    "h3": ("Heading 2", Pt(14), Pt(12), Pt(5),  False),
    "h4": ("Heading 3", Pt(12), Pt(8),  Pt(3),  False),
}

def add_heading(doc, element):
    tag = element.name
    style_name, font_sz, sp_before, sp_after, add_rule = _HEADING_MAP[tag]

    try:
        p = doc.add_paragraph(style=style_name)
    except Exception:
        p = doc.add_paragraph()
    p.clear()
    add_inline_runs(p, element, bold=True, size=font_sz)

    for run in p.runs:
        run.font.name  = FONT_NAME
        run.font.size  = font_sz
        run.font.color.rgb = BRAND_DARK

    p.paragraph_format.space_before = sp_before
    p.paragraph_format.space_after  = sp_after

    strip_para_ends(p)

    if add_rule:
        add_bottom_border_to_para(p, "3498DB", sz="4")


# ── Element dispatcher ────────────────────────────────────────────────────────
def process_element(doc, element):
    if not hasattr(element, "name") or element.name is None:
        return

    tag = element.name

    # ── Skip the cover-page items (already rendered manually) ──────────────
    if tag == "h1":
        return  # title is on the cover page

    if tag == "p":
        style_attr = element.get("style", "")
        # The intro subtitle paragraph has an explicit font-size inline style
        if "font-size" in style_attr and "text-align" in style_attr:
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(2)
        add_inline_runs(p, element, size=Pt(11))
        strip_para_ends(p)
        return

    # ── Structural elements ────────────────────────────────────────────────
    if tag in ("h2", "h3", "h4"):
        add_heading(doc, element)

    elif tag == "pre":
        add_code_block(doc, element.get_text())

    elif tag in ("ul", "ol"):
        add_list_items(doc, element, ordered=(tag == "ol"))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    elif tag == "table":
        add_html_table(doc, element)

    elif tag == "img":
        add_image_element(doc, element)

    elif tag == "div":
        cls_list = element.get("class", [])
        cls_str  = " ".join(cls_list) if isinstance(cls_list, list) else cls_list

        if "toc" in cls_str:
            return  # skip in-HTML TOC — Word TOC is on the cover page

        if "flowchart-container" in cls_str:
            h4_el = element.find("h4")
            title = h4_el.get_text(strip=True) if h4_el else "Application Data Flow & Architecture"
            if h4_el:
                add_heading(doc, h4_el)
            
            # Use a mock object matching the BeautifulSoup tag interface for add_image_element
            class MockImgTag:
                def __init__(self, src, alt):
                    self.src = src
                    self.alt = alt
                def get(self, key, default=None):
                    if key == "src":
                        return self.src
                    if key == "alt":
                        return self.alt
                    return default

            add_image_element(doc, MockImgTag("images/flowchart.png", title))
            return

        add_callout(doc, element)  # handles note/warning/success; falls back to recurse

    elif tag == "hr":
        hr = doc.add_paragraph()
        add_bottom_border_to_para(hr, "CCCCCC", sz="4")
        hr.paragraph_format.space_after = Pt(4)


# ── Header & Footer (matching reference document style) ───────────────────────
LOGO_PATH = BASE_DIR / "branding" / "logo_sq.png"


def _add_page_field(para):
    """Insert a { PAGE } auto-number field into *para*."""
    rpr_attribs = {"w:color": "000000", "w:sz": "22", "w:szCs": "22"}

    def _run_with_rpr(fld_type=None, text=None):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for k, v in rpr_attribs.items():
            el = OxmlElement(k)
            el.set(qn(k), v)
            rPr.append(el)
        r.append(rPr)
        if fld_type:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type)
            r.append(fc)
        if text is not None:
            t = OxmlElement("w:t")
            t.text = text
            r.append(t)
        para._p.append(r)

    _run_with_rpr(fld_type="begin")
    # instrText run
    r_instr = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    for k, v in rpr_attribs.items():
        el = OxmlElement(k)
        el.set(qn(k), v)
        rPr2.append(el)
    r_instr.append(rPr2)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    r_instr.append(instr)
    para._p.append(r_instr)

    _run_with_rpr(fld_type="separate")
    _run_with_rpr(text="1")   # placeholder — Word replaces this
    _run_with_rpr(fld_type="end")


def build_header_footer(doc, version_info: dict):
    """
    Update the template's existing headers and footers in-place,
    substituting Decibels Lab references with Bytehound branding.
    This preserves the exact design, formatting, and embedded shapes/lines.
    """
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    version = version_info.get("version", "1.0.0")

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Update Header text nodes in-place
    for t in sec.header._element.findall('.//w:t', ns):
        if t.text == "Decibels Lab Private Limited":
            t.text = "Bytehound Open Source Community"
        elif t.text == "CIN: U80904KA2019PTC126675":
            t.text = "Serial Telemetry & Decoder Guide"
        elif t.text == "+":
            t.text = "github.com/P-Shreyas543/Bytehound"
        elif t.text and ("91 89515 03455" in t.text or "contact@decibelslab.com" in t.text):
            t.text = ""

    # Update Footer text nodes in-place
    for t in sec.footer._element.findall('.//w:t', ns):
        if t.text == "lms.decibelslab.com":
            t.text = "github.com/P-Shreyas543/Bytehound"
        elif t.text and "Multi-Cell BMS Algorithm Development Kit" in t.text:
            t.text = f"  /  Bytehound v{version} \u2014 Serial Telemetry & Decoder Guide"


# ── Body-content wipe (preserves headers / footers / section properties) ──────
def clear_body_content(doc):
    """
    Remove every w:p and w:tbl from the document body while keeping
    w:sectPr (section properties), which holds the header/footer references.
    Without sectPr the header and footer links would be lost.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    version_info: dict = {}
    if VERSION_PATH.exists():
        with open(VERSION_PATH, encoding="utf-8") as f:
            version_info = json.load(f)

    print(f"Bytehound User Manual generator  (v{version_info.get('version', '?')})")
    print(f"  Template : {TEMPLATE_PATH}")
    print(f"  Source   : {HTML_PATH}")
    print(f"  Output   : {OUT_PATH}")

    print("\nOpening template document...")
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))

    print("Clearing body content (keeping header & footer)...")
    clear_body_content(doc)

    print("Building cover page...")
    build_cover_page(doc, version_info)

    print("Building headers and footers...")
    build_header_footer(doc, version_info)

    print("Parsing HTML...")
    with open(HTML_PATH, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    container = soup.select_one(".content-body") or soup.select_one("div.container")
    if container is None:
        container = soup.find("body")
        print("  Warning: Content container not found — falling back to <body>.")

    children = [el for el in container.children
                if hasattr(el, "name") and el.name]
    print(f"  Found {len(children)} top-level elements.")

    for element in children:
        process_element(doc, element)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"\nSaved: {OUT_PATH}")
    print("Tip: open in Word and press Ctrl+A then F9 to populate the Table of Contents.")


if __name__ == "__main__":
    main()
